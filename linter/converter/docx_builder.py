import copy
import re
from pathlib import Path
import mistune
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from linter.config import Config
from linter.converter.tables.processor import (
    apply_table_preset,
    apply_zebra_preset,
    fit_table_to_page,
    estimate_column_weights,
    normalize_width_weights,
    format_cell_paragraph,
    keep_table_together,
    get_section_column_width,
)


def _force_auto_height_and_cols(section, count, space=420, equal_width=True):
    sectPr = section._sectPr
    for va in sectPr.xpath('./w:vAlign'):
        sectPr.remove(va)
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.insert(0, cols)
    cols.set(qn('w:num'), str(count))
    cols.set(qn('w:space'), str(space))
    cols.set(qn('w:equalWidth'), '1' if equal_width else '0')


def _clear_document_content(doc):
    body = doc.element.body
    for p in body.xpath('./w:p'):
        body.remove(p)
    for tbl in body.xpath('./w:tbl'):
        body.remove(tbl)
    sections = list(doc.sections)
    for section in sections[1:]:
        body.remove(section._element)


def _disable_contextual_spacing(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    cs = pPr.find(qn('w:contextualSpacing'))
    if cs is not None:
        pPr.remove(cs)


def _add_space_after(paragraph, points=25):
    _disable_contextual_spacing(paragraph)
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), str(int(points * 20)))


def _set_space_before(paragraph, points=0):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(points * 20)))


def _set_space_after(paragraph, points=0):
    _disable_contextual_spacing(paragraph)
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), str(int(points * 20)))


def _add_spacing_paragraph(doc, points=25):
    spacer = doc.add_paragraph("\u00a0")
    _add_space_after(spacer, points)
    return spacer


def _get_text_recursive(node):
    if 'raw' in node:
        return node['raw']
    if 'text' in node:
        return node['text']
    if 'children' in node:
        return "".join(_get_text_recursive(c) for c in node['children'])
    return ""


_FONT_COLOR_RE = re.compile(r'<font\s+color=["\']#?([0-9a-fA-F]{6})["\'].*?>')
_FONT_CLOSE_RE = re.compile(r'</font\s*>')
_BREAK_HTML_RE = re.compile(r'<!--break:(page|column)-->')
_DEFAULT_HYPERLINK_STYLE = 'Hyperlink'


def _get_or_create_hyperlink_rel(part, url):
    for rel in part.rels.values():
        if rel.is_external and rel.target_ref == url:
            return rel.rId
    return part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )


def _append_run_style(rPr, style_name: str = None):
    if not style_name:
        return
    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), style_name)
    rPr.append(r_style)


def _append_underline(rPr, underline_value: str = 'single'):
    u = OxmlElement('w:u')
    u.set(qn('w:val'), underline_value)
    rPr.append(u)


def _build_hyperlink_runs(parent_elem, children, bold=False, italic=False, color_stack=None,
                          link_style: str = _DEFAULT_HYPERLINK_STYLE):
    if color_stack is None:
        color_stack = []

    for child in children:
        c_type = child.get('type')
        raw_text = child.get('raw') or child.get('text', '')

        if c_type == 'text' and raw_text:
            run_elem = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            _append_run_style(rPr, link_style)

            if bold:
                rPr.append(OxmlElement('w:b'))
            if italic:
                rPr.append(OxmlElement('w:i'))

            if not link_style:
                _append_underline(rPr)

            if color_stack:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), color_stack[-1])
                rPr.append(c)

            run_elem.append(rPr)
            t = OxmlElement('w:t')
            t.text = raw_text
            run_elem.append(t)
            parent_elem.append(run_elem)

        elif c_type in ('strong', 'double_emphasis'):
            _build_hyperlink_runs(
                parent_elem, child.get('children', [child]),
                bold=True, italic=italic, color_stack=color_stack, link_style=link_style
            )

        elif c_type in ('emphasis', 'italic'):
            _build_hyperlink_runs(
                parent_elem, child.get('children', [child]),
                bold=bold, italic=True, color_stack=color_stack, link_style=link_style
            )

        elif c_type == 'inline_html':
            m = _FONT_COLOR_RE.match(raw_text)
            if m:
                color_stack.append(m.group(1))
                continue
            if _FONT_CLOSE_RE.match(raw_text):
                if color_stack:
                    color_stack.pop()
                continue
            if raw_text:
                run_elem = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                _append_run_style(rPr, link_style)
                if bold:
                    rPr.append(OxmlElement('w:b'))
                if italic:
                    rPr.append(OxmlElement('w:i'))
                if not link_style:
                    _append_underline(rPr)
                run_elem.append(rPr)
                t = OxmlElement('w:t')
                t.text = raw_text
                run_elem.append(t)
                parent_elem.append(run_elem)

        elif 'children' in child:
            _build_hyperlink_runs(
                parent_elem, child['children'],
                bold=bold, italic=italic, color_stack=color_stack, link_style=link_style
            )


def _add_hyperlink_with_children(paragraph, part, url, children, bold=False, italic=False,
                                 color_stack=None, link_style: str = _DEFAULT_HYPERLINK_STYLE):
    if not url:
        return
    rId = _get_or_create_hyperlink_rel(part, url)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rId)
    _build_hyperlink_runs(
        hyperlink, children, bold=bold, italic=italic,
        color_stack=color_stack or [], link_style=link_style
    )
    paragraph._p.append(hyperlink)


def _has_link(node):
    if node.get('type') == 'link':
        return True
    for child in node.get('children', []):
        if _has_link(child):
            return True
    return False


def _children_have_link(children):
    for child in children:
        if _has_link(child):
            return True
    return False


def _resolve_character_style_id(doc, style_name: str):
    if not style_name:
        return None
    for style in doc.styles:
        if style.type == 2 and (style.name == style_name or style.style_id == style_name):
            return style.style_id
    return style_name


def _add_formatted_text(paragraph, children, bold=False, italic=False, underline=False,
                        font_color=None, color_stack=None, part=None,
                        link_style: str = _DEFAULT_HYPERLINK_STYLE):
    if color_stack is None:
        color_stack = []

    for child in children:
        c_type = child.get('type')
        raw_text = child.get('raw') or child.get('text', '')

        if c_type == 'inline_html':
            break_m = _BREAK_HTML_RE.match(raw_text)
            if break_m:
                run = paragraph.add_run()
                if break_m.group(1) == 'page':
                    run.add_break(WD_BREAK.PAGE)
                else:
                    run.add_break(WD_BREAK.COLUMN)
                continue

            m = _FONT_COLOR_RE.match(raw_text)
            if m:
                color_stack.append(m.group(1))
                continue
            if _FONT_CLOSE_RE.match(raw_text):
                if color_stack:
                    color_stack.pop()
                continue
            if raw_text:
                run = paragraph.add_run(raw_text)
                run.bold = bold
                run.italic = italic
                run.underline = underline

        elif c_type == 'text':
            if raw_text:
                run = paragraph.add_run(raw_text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                if color_stack:
                    run.font.color.rgb = RGBColor.from_string(color_stack[-1])
                elif font_color:
                    run.font.color.rgb = font_color

        elif c_type in ('strong', 'double_emphasis'):
            if 'children' in child:
                _add_formatted_text(paragraph, child['children'], bold=True, italic=italic,
                                    underline=underline, color_stack=color_stack, part=part)
            else:
                text = _get_text_recursive(child)
                if text:
                    run = paragraph.add_run(text)
                    run.bold = True
                    run.italic = italic
                    run.underline = underline
                    if color_stack:
                        run.font.color.rgb = RGBColor.from_string(color_stack[-1])
                    elif font_color:
                        run.font.color.rgb = font_color

        elif c_type in ('emphasis', 'italic'):
            if 'children' in child:
                _add_formatted_text(paragraph, child['children'], bold=bold, italic=True,
                                    underline=underline, color_stack=color_stack, part=part)
            else:
                text = _get_text_recursive(child)
                if text:
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = True
                    run.underline = underline
                    if color_stack:
                        run.font.color.rgb = RGBColor.from_string(color_stack[-1])
                    elif font_color:
                        run.font.color.rgb = font_color

        elif c_type == 'link':
            url = child.get('attrs', {}).get('url', '')
            if url and part:
                _add_hyperlink_with_children(
                    paragraph, part, url,
                    child.get('children', []),
                    bold=bold, italic=italic,
                    color_stack=list(color_stack),
                    link_style=link_style
                )
            elif 'children' in child:
                _add_formatted_text(paragraph, child['children'], bold=bold, italic=italic,
                                    underline=True, color_stack=color_stack, part=part,
                                    link_style=link_style)
            else:
                text = _get_text_recursive(child)
                if text:
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    if link_style:
                        try:
                            run.style = link_style
                        except Exception:
                            pass
                    else:
                        run.underline = True
                    if color_stack:
                        run.font.color.rgb = RGBColor.from_string(color_stack[-1])
                    elif font_color:
                        run.font.color.rgb = font_color

        elif 'children' in child:
            _add_formatted_text(paragraph, child['children'], bold=bold, italic=italic,
                                underline=underline, color_stack=color_stack, part=part,
                                link_style=link_style)


def _apply_heading_style_in_cell(paragraph):
    if not paragraph.runs:
        return
    first_text = paragraph.runs[0].text
    if not first_text.startswith('#'):
        return
    level = 0
    for ch in first_text:
        if ch == '#':
            level += 1
        else:
            break
    if not (1 <= level <= 6):
        return
    rest_start = level
    while rest_start < len(first_text) and first_text[rest_start] == ' ':
        rest_start += 1
    rest = first_text[rest_start:]
    paragraph.runs[0].text = rest
    try:
        paragraph.style = f"Heading {level}"
    except Exception:
        for run in paragraph.runs:
            if run.text:
                run.bold = True


_INDENT_BLOCK_ALIASES = {
    'horizontal rule': 'thematic_break',
    'horizontal line': 'thematic_break',
    'горизонтальная линия': 'thematic_break',
}


class DocxBuilder:

    def __init__(self, template_path: Path, config: Config):
        self.template_path = template_path
        self.config = config
        self.doc = None
        self._current_block_style = None
        self._current_block_style_config = None
        self._last_block_type = None
        self._last_para_style_name = None
        self._last_paragraph = None
        self._is_first_para_in_block = False
        self._suppress_next_first_line_indent = False
        self._suppress_next_first_line_indent_from_arrow = False
        self._hyperlink_style_id = _DEFAULT_HYPERLINK_STYLE
        self._cols_num = 1
        self._cols_space = 420
        self._cols_equal_width = True
        self._markdown_parser = None

    def build(self, processed_md: str) -> None:
        self.doc = Document(str(self.template_path))
        _clear_document_content(self.doc)
        self._hyperlink_style_id = _resolve_character_style_id(self.doc, _DEFAULT_HYPERLINK_STYLE)

        if self.doc.sections:
            sect_pr = self.doc.sections[0]._sectPr
            template_cols = sect_pr.find(qn('w:cols'))
            if template_cols is not None:
                self._cols_num = int(template_cols.get(qn('w:num'), 2))
                self._cols_space = int(template_cols.get(qn('w:space'), 420))
                eq = template_cols.get(qn('w:equalWidth'), '1')
                self._cols_equal_width = eq == '1'
            _force_auto_height_and_cols(self.doc.sections[0], self._cols_num, self._cols_space, self._cols_equal_width)

        print("[*] Верстаю документ...")

        self._markdown_parser = mistune.create_markdown(renderer=None, plugins=['table'])
        ast = self._markdown_parser(processed_md)
        self._process_ast(ast)
        self._apply_table_spacing()

    def save(self, output_path: Path) -> None:
        if self.doc is None:
            raise RuntimeError("Документ не построен. Сначала вызовите build().")
        self.doc.save(str(output_path))

    def _process_table(self, table_node: dict, table_config: dict = None) -> None:
        if table_config is None:
            table_config = self.config.table_mappings.get("‹!plain›", {
                "table": "Простая таблица 2",
                "header": "Table Header 2",
                "text": "Table Text 2",
                "preset": "plain"
            })

        preset_name = table_config.get('preset', 'plain')

        if not table_node.get('children') or len(table_node['children']) < 1:
            return

        header_row_node = table_node['children'][0]
        header_cells_nodes = header_row_node.get('children', [])

        body_rows_nodes = []
        if len(table_node['children']) > 1:
            tbody_node = table_node['children'][1]
            body_rows_nodes = tbody_node.get('children', [])

        cols_count = len(header_cells_nodes)
        if cols_count == 0:
            return

        table = self.doc.add_table(rows=0, cols=cols_count)

        preset = self.config.table_presets.get(preset_name, self.config.table_presets.get('plain', {}))

        if preset.get('renderer') != 'zebra':
            try:
                table.style = table_config.get('table', 'Простая таблица 2')
            except Exception:
                try:
                    table.style = 'Сетка таблицы'
                except Exception:
                    pass

        header_row = table.add_row()
        for idx, cell_node in enumerate(header_cells_nodes):
            if idx < cols_count:
                p = header_row.cells[idx].paragraphs[0]
                cell_children = cell_node.get('children', [])
                if cell_children:
                    _add_formatted_text(p, cell_children, part=self.doc.part, link_style=self._hyperlink_style_id)
                format_cell_paragraph(p, table_config.get('header'), bold=False)
                _apply_heading_style_in_cell(p)

        for row_node in body_rows_nodes:
            current_row = table.add_row()
            row_cells_nodes = row_node.get('children', [])
            for idx, cell_node in enumerate(row_cells_nodes):
                if idx < cols_count:
                    p = current_row.cells[idx].paragraphs[0]
                    cell_children = cell_node.get('children', [])
                    if cell_children:
                        _add_formatted_text(p, cell_children, part=self.doc.part, link_style=self._hyperlink_style_id)
                    format_cell_paragraph(p, table_config.get('text'))
                    _apply_heading_style_in_cell(p)

        width_weights = normalize_width_weights(
            preset.get('column_widths') or table_config.get('column_widths'),
            cols_count
        )
        if width_weights is None:
            def text_getter(node):
                return _get_text_recursive(node)
            width_weights = estimate_column_weights(
                header_cells_nodes, body_rows_nodes, cols_count, text_getter
            )

        fit_to_page = preset.get('fit_to_page', True) if isinstance(preset.get('fit_to_page'), bool) else True
        row_heights = preset.get('row_heights') if isinstance(preset.get('row_heights'), list) else None
        table_width = preset.get('table_width') if isinstance(preset.get('table_width'), (int, float)) else None

        apply_table_preset(table, preset_name, self.config.table_presets)

        fit_table_to_page(table, self.doc.sections[-1], cols_count, width_weights, fit_to_page, row_heights, table_width)

        empty_row_top = preset.get('empty_row_top', False)
        empty_row_bottom = preset.get('empty_row_bottom', False)
        empty_row_height = preset.get('empty_row_height', 200)

        if empty_row_bottom:
            bottom_row = table.add_row()
            self._make_empty_row(bottom_row, empty_row_height)

        if empty_row_top:
            top_row = table.add_row()
            self._make_empty_row(top_row, empty_row_height)
            tr_elements = table._tbl.findall(qn('w:tr'))
            if len(tr_elements) > 1:
                tr = tr_elements[-1]
                table._tbl.remove(tr)
                first_tr = tr_elements[0]
                table._tbl.insert(table._tbl.index(first_tr), tr)

        keep_table_together(table)

    @staticmethod
    def _make_empty_row(row, height_twips):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.clear()
            tcPr = cell._tc.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                shd = OxmlElement('w:shd')
                tcPr.append(shd)
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'auto')
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for edge in ('top', 'left', 'bottom', 'right'):
                edge_el = tcBorders.find(qn(f'w:{edge}'))
                if edge_el is None:
                    edge_el = OxmlElement(f'w:{edge}')
                    tcBorders.append(edge_el)
                edge_el.set(qn('w:val'), 'nil')

        trPr = row._tr.get_or_add_trPr()
        trHeight = trPr.find(qn('w:trHeight'))
        if trHeight is None:
            trHeight = OxmlElement('w:trHeight')
            trPr.append(trHeight)
        trHeight.set(qn('w:val'), str(height_twips))
        trHeight.set(qn('w:hRule'), 'exact')

    def _apply_table_spacing(self) -> None:
        from docx.text.paragraph import Paragraph
        body = self.doc.element.body
        prev = None
        for child in list(body):
            tag = child.tag
            if prev is not None and tag in (qn('w:p'), qn('w:tbl')):
                if prev.tag == qn('w:tbl'):
                    if tag == qn('w:p'):
                        para = Paragraph(child, self.doc)
                        is_heading = para.style is not None and para.style.name.startswith('Heading')
                        points = self.config.table_before_heading if is_heading else self.config.after_table_spacing
                        if points <= 0:
                            prev = child
                            continue
                        pPr = child.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            child.insert(0, pPr)
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        cur = int(spacing.get(qn('w:before'), '0'))
                        twips = int(points * 20)
                        if cur < twips:
                            spacing.set(qn('w:before'), str(twips))
                    elif tag == qn('w:tbl'):
                        after = self.config.after_table_spacing
                        if after <= 0:
                            prev = child
                            continue
                        tblPr = child.find(qn('w:tblPr'))
                        if tblPr is None:
                            tblPr = OxmlElement('w:tblPr')
                            child.insert(0, tblPr)
                        tblSpacing = tblPr.find(qn('w:tblSpacing'))
                        if tblSpacing is None:
                            tblSpacing = OxmlElement('w:tblSpacing')
                            tblPr.append(tblSpacing)
                        cur = int(tblSpacing.get(qn('w:before'), '0'))
                        twips = int(after * 20)
                        if cur < twips:
                            tblSpacing.set(qn('w:before'), str(twips))
            prev = child

    def _track_keep_together_para(self, para):
        if self._in_keep_together and para is not None:
            para.paragraph_format.keep_together = True
            if self._keep_together_last_para is not None:
                self._keep_together_last_para.paragraph_format.keep_with_next = True
            self._keep_together_last_para = para

    def _process_ast(self, ast: list) -> None:
        last_paragraph = None
        has_content = False
        next_table_config = None
        self._in_keep_together = False
        self._keep_together_last_para = None

        i = 0
        n = len(ast)

        while i < n:
            node = ast[i]
            n_type = node.get('type')

            if n_type == 'blank_line':
                i += 1
                continue

            if n_type == 'paragraph':
                full_text = _get_text_recursive(node)
                check_text = full_text.strip().lower()

                if check_text in self.config.table_mappings:
                    next_table_config = self.config.table_mappings[check_text]
                    print(f"[i] Найдена метка таблицы: {check_text} -> preset: {next_table_config.get('preset', 'plain')}")
                    i += 1
                    continue

            full_text = _get_text_recursive(node)
            if 'raw' in node:
                full_text = node['raw']

            check_text = full_text.strip()
            if check_text.startswith("<!--block_start:") or check_text.startswith("<!--single_tag:"):
                marker_para = self._handle_style_marker(node, check_text)
                if marker_para is not None:
                    last_paragraph = marker_para
                    self._track_keep_together_para(last_paragraph)
                    next_is_heading_flag = False
                    for jj in range(i + 1, n):
                        nn = ast[jj]
                        if nn.get('type') == 'blank_line':
                            continue
                        if nn.get('type') in ('paragraph', 'block_html'):
                            t = _get_text_recursive(nn)
                            if t.strip().startswith('<!--'):
                                continue
                        if nn.get('type') == 'heading':
                            nl = nn.get('attrs', {}).get('level', 1)
                            if nl <= 3:
                                next_is_heading_flag = True
                        break
                    if next_is_heading_flag:
                        _add_space_after(last_paragraph, self.config.before_heading_spacing)
                i += 1
                continue
            if check_text.startswith("<!--block_end:"):
                self._handle_style_marker(node, check_text)
                i += 1
                continue
            if check_text.startswith("<!--break:"):
                self._handle_break_marker(check_text)
                last_paragraph = None
                self._keep_together_last_para = None
                has_content = True
                self._last_block_type = None
                self._last_para_style_name = None
                self._last_paragraph = None
                self._suppress_next_first_line_indent = True
                i += 1
                continue

            if check_text == "<!--keep_together_start-->":
                self._in_keep_together = True
                i += 1
                continue
            if check_text == "<!--keep_together_end-->":
                self._in_keep_together = False
                self._keep_together_last_para = None
                i += 1
                continue

            if check_text.startswith("<!--line_style:"):
                style_key = self._get_line_style_key(check_text)
                style_name = self._get_line_style_name(style_key)
                last_paragraph = self.doc.add_paragraph("\u00a0", style=style_name)
                self._track_keep_together_para(last_paragraph)
                has_content = True
                self._last_block_type = 'thematic_break'
                self._last_para_style_name = style_name
                self._last_paragraph = last_paragraph

                arrows = self._get_line_style_arrows(check_text)
                if arrows:
                    space_count = arrows.count('›')
                    if space_count:
                        _add_space_after(last_paragraph, space_count * 5)
                    if '→' in arrows:
                        self._suppress_next_first_line_indent_from_arrow = True
                    filler_count = arrows.count('⋙')
                    for _ in range(filler_count):
                        spacer = self.doc.add_paragraph()
                        self._track_keep_together_para(spacer)
                        try:
                            spacer.style = last_paragraph.style.name
                        except Exception:
                            pass
                        _set_space_after(spacer, 0)
                        _set_space_before(spacer, 0)
                        _disable_contextual_spacing(spacer)

                for jj in range(i + 1, n):
                    nn = ast[jj]
                    if nn.get('type') == 'blank_line':
                        continue
                    if nn.get('type') in ('paragraph', 'block_html'):
                        t = _get_text_recursive(nn)
                        if t.strip().startswith('<!--'):
                            continue
                    if nn.get('type') == 'heading':
                        nl = nn.get('attrs', {}).get('level', 1)
                        if nl <= 3:
                            _add_space_after(last_paragraph, self.config.before_heading_spacing)
                    break
                i += 1
                continue

            current_block_type = 'paragraph'
            if n_type == 'heading':
                current_block_type = 'heading'
            elif n_type == 'list':
                current_block_type = 'list'
            elif n_type == 'thematic_break':
                current_block_type = 'thematic_break'
            elif n_type == 'table':
                current_block_type = 'table'

            next_is_heading = False
            next_is_list = False
            next_is_table = False
            next_node_type = None

            for j in range(i + 1, n):
                next_node = ast[j]
                if next_node.get('type') == 'blank_line':
                    continue
                if next_node.get('type') in ('paragraph', 'block_html'):
                    txt = _get_text_recursive(next_node)
                    stripped = txt.strip()
                    if stripped.startswith('<!--'):
                        if stripped.startswith('<!--line_style:'):
                            break
                        continue
                    if stripped.lower() in self.config.table_mappings:
                        continue
                next_node_type = next_node.get('type')
                if next_node_type == 'heading':
                    next_level = next_node.get('attrs', {}).get('level', 1)
                    if next_level <= 3:
                        next_is_heading = True
                if next_node_type == 'list':
                    next_is_list = True
                if next_node_type == 'table':
                    next_is_table = True
                break

            if n_type == 'table':
                table_conf = next_table_config
                next_table_config = None
                self._process_table(node, table_conf)
                last_paragraph = None
                has_content = True
                self._last_block_type = 'table'
                self._last_para_style_name = None
                self._last_paragraph = None
                self._suppress_next_first_line_indent_from_arrow = False

            elif n_type == 'heading':
                self._current_block_style = None
                self._current_block_style_config = None
                last_paragraph = self._process_heading(node, has_content)
                has_content = True
                self._last_block_type = 'heading'
                self._last_para_style_name = None
                self._last_paragraph = None
                self._suppress_next_first_line_indent_from_arrow = False

            elif n_type == 'list':
                is_ordered = node.get('attrs', {}).get('ordered', False)
                last_paragraph = self._process_list(node, is_ordered)
                has_content = True
                self._last_block_type = 'list'
                self._last_para_style_name = None
                self._suppress_next_first_line_indent_from_arrow = False

            elif n_type == 'thematic_break':
                last_paragraph = self.doc.add_paragraph("\u00a0", style="Horizontal Line")
                has_content = True
                self._last_block_type = 'thematic_break'
                self._last_para_style_name = 'Horizontal Line'
                self._last_paragraph = last_paragraph

            else:
                style_for_para = self._current_block_style
                last_paragraph = self._process_paragraph(node, style_for_para)
                has_content = True
                self._last_block_type = 'paragraph'
                self._is_first_para_in_block = False

            if next_is_heading and last_paragraph and current_block_type != 'heading':
                _add_space_after(last_paragraph, self.config.before_heading_spacing)

            if next_is_table and last_paragraph:
                _add_space_after(last_paragraph, self.config.before_table_spacing)

            if (current_block_type == 'list' and
                last_paragraph and
                not next_is_list and
                not next_is_heading and
                not next_is_table and
                next_node_type is not None):
                _add_space_after(last_paragraph, self.config.after_list_spacing)

            if n_type != 'table':
                self._track_keep_together_para(last_paragraph)
            i += 1

    def _process_heading(self, node: dict, has_content: bool):
        level = node.get('attrs', {}).get('level', 1)
        children = node.get('children', [])

        if level == 1:
            if not has_content:
                _force_auto_height_and_cols(self.doc.sections[0], 1)
            else:
                sec_h1 = self.doc.add_section(WD_SECTION.NEW_PAGE)
                _force_auto_height_and_cols(sec_h1, 1)

            h = self.doc.add_paragraph(style='Heading 1')
            _add_formatted_text(h, children, part=self.doc.part, link_style=self._hyperlink_style_id)
            h.alignment = 1

            sec_text = self.doc.add_section(WD_SECTION.CONTINUOUS)
            _force_auto_height_and_cols(sec_text, self._cols_num, self._cols_space, self._cols_equal_width)
            return h
        else:
            h = self.doc.add_paragraph(style=f'Heading {level}')
            _add_formatted_text(h, children, part=self.doc.part, link_style=self._hyperlink_style_id)
            return h

    def _process_list(self, list_node: dict, ordered: bool = False, level: int = 0):
        style_name = self._resolve_list_style_name(ordered, level)

        last_p = None
        for item in list_node.get('children', []):
            if item.get('type') != 'list_item':
                continue

            p = self.doc.add_paragraph(style=style_name)
            self._fix_list_paragraph(p, level, style_name)
            self._flatten_style_properties(p)
            self._collapse_spacing_for_related_block_styles(p, style_name)
            last_p = p

            inline_nodes = []
            for child in item.get('children', []):
                if child.get('type') == 'list':
                    is_ordered_nested = child.get('attrs', {}).get('ordered', False)
                    nested_last = self._process_list(child, is_ordered_nested, level + 1)
                    if nested_last:
                        last_p = nested_last
                else:
                    inline_nodes.append(child)

            if inline_nodes:
                    _add_formatted_text(p, inline_nodes, part=self.doc.part, link_style=self._hyperlink_style_id)
                    self._apply_indent_from_arrows(p)
                    if self._suppress_next_first_line_indent_from_arrow:
                        self._suppress_next_first_line_indent_from_arrow = False

            self._last_para_style_name = style_name
            self._last_paragraph = p

        return last_p

    def _flatten_style_properties(self, paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        style = paragraph.style
        if style is None:
            return
        style_pPr = style.element.find(qn('w:pPr'))
        if style_pPr is None:
            return

        for style_child in list(style_pPr):
            tag = style_child.tag
            if tag == qn('w:pStyle'):
                continue
            if pPr.find(tag) is None:
                pPr.append(copy.deepcopy(style_child))

    def _fix_list_paragraph(self, p, level=0, expected_style=""):
        style = p.style
        pPr = p._element.get_or_add_pPr()

        if pPr.find(qn('w:numPr')) is None:
            style_pPr = style.element.find(qn('w:pPr'))
            if style_pPr is not None:
                style_numPr = style_pPr.find(qn('w:numPr'))
                if style_numPr is not None:
                    new_numPr = copy.deepcopy(style_numPr)
                    is_fallback = style.name != expected_style
                    if is_fallback:
                        ilvl = new_numPr.find(qn('w:ilvl'))
                        if ilvl is not None:
                            ilvl.set(qn('w:val'), str(level))
                    pPr.append(new_numPr)

    def _apply_indent_from_arrows(self, paragraph):
        runs = paragraph.runs
        if not runs:
            return
        last_run = runs[-1]
        text = last_run.text
        if not text:
            return
        m = re.search(r'([›→⋙]+)\s*$', text)
        if m:
            markers = m.group(1)
            text = text[:m.start()]
            last_run.text = text
            space_count = markers.count('›')
            if space_count:
                _add_space_after(paragraph, space_count * 5)
            if '→' in markers:
                self._suppress_next_first_line_indent_from_arrow = True
            filler_count = markers.count('⋙')
            for _ in range(filler_count):
                spacer = self.doc.add_paragraph()
                try:
                    spacer.style = paragraph.style.name
                except Exception:
                    pass
                _set_space_after(spacer, 0)
                _set_space_before(spacer, 0)
                _disable_contextual_spacing(spacer)

    def _apply_pending_break_indent_rule(self, paragraph):
        if not self._suppress_next_first_line_indent:
            return
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        self._suppress_next_first_line_indent = False

    def _build_paragraph_from_children(self, children, style_name: str = None):
        if not children:
            return None

        p = self.doc.add_paragraph()
        if style_name:
            try:
                p.style = style_name
                self._flatten_style_properties(p)
            except Exception:
                pass

        _add_formatted_text(p, children, part=self.doc.part, link_style=self._hyperlink_style_id)

        if style_name and self._is_first_para_in_block and self.config.no_indent_after_block_types:
            if style_name in self.config.no_indent_after_block_types:
                p.paragraph_format.first_line_indent = Pt(0)

        if style_name is None and self.config.no_indent_after_block_types:
            remove = False
            for t in self.config.no_indent_after_block_types:
                if t == self._last_block_type or t == self._last_para_style_name:
                    remove = True
                    break
                alias = _INDENT_BLOCK_ALIASES.get(t)
                if alias and (alias == self._last_block_type or alias == self._last_para_style_name):
                    remove = True
                    break
            if not remove and self.config.no_indent_if_first_bold:
                if p.runs and p.runs[0].bold:
                    remove = True
            if remove:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Pt(0)

        self._apply_pending_break_indent_rule(p)
        if self._suppress_next_first_line_indent_from_arrow:
            p.paragraph_format.first_line_indent = Pt(0)
            self._suppress_next_first_line_indent_from_arrow = False
        self._collapse_spacing_for_related_block_styles(p, style_name)
        self._last_para_style_name = style_name
        self._last_paragraph = p
        self._apply_indent_from_arrows(p)
        return p

    def _style_matches_list_base(self, style_name: str, base_style: str) -> bool:
        if not style_name or not base_style:
            return False
        return style_name == base_style or style_name.startswith(base_style + " ")

    def _collapse_spacing_for_related_block_styles(self, paragraph, style_name: str):
        if not self._current_block_style_config or not style_name or self._last_paragraph is None:
            return

        block_style = (self._current_block_style_config.style_name or "").strip()
        prev_style = (self._last_para_style_name or "").strip()
        curr_style = style_name.strip()
        bullet_base = (self._current_block_style_config.list_bullet_style or "").strip()
        number_base = (self._current_block_style_config.list_number_style or "").strip()

        prev_is_block = prev_style == block_style
        curr_is_block = curr_style == block_style
        prev_is_list = (
            self._style_matches_list_base(prev_style, bullet_base) or
            self._style_matches_list_base(prev_style, number_base)
        )
        curr_is_list = (
            self._style_matches_list_base(curr_style, bullet_base) or
            self._style_matches_list_base(curr_style, number_base)
        )

        if (prev_is_block and curr_is_list) or (prev_is_list and curr_is_block):
            _set_space_after(self._last_paragraph, 0)
            _set_space_before(paragraph, 0)

    def _resolve_list_style_base(self, ordered: bool) -> str:
        if self._current_block_style_config:
            candidate = (
                self._current_block_style_config.list_number_style
                if ordered else
                self._current_block_style_config.list_bullet_style
            )
            if candidate:
                return candidate
        return "List Number" if ordered else "List Bullet"

    def _resolve_list_style_name(self, ordered: bool, level: int) -> str:
        default_base = "List Number" if ordered else "List Bullet"
        preferred_base = self._resolve_list_style_base(ordered)
        available = {s.name for s in self.doc.styles}

        candidates = []
        if level == 0:
            candidates.append(preferred_base)
            if preferred_base != default_base:
                candidates.append(default_base)
        else:
            candidates.append(f"{preferred_base} {level + 1}")
            candidates.append(preferred_base)
            if preferred_base != default_base:
                candidates.append(f"{default_base} {level + 1}")
                candidates.append(default_base)

        for candidate in candidates:
            if candidate in available:
                return candidate

        return default_base

    def _process_paragraph(self, node: dict, style_name: str = None):
        children = node.get('children', [])
        if not children:
            return None

        if not any(
            child.get('type') == 'inline_html' and _BREAK_HTML_RE.match(child.get('raw') or child.get('text', ''))
            for child in children
        ):
            return self._build_paragraph_from_children(children, style_name)

        last_paragraph = None
        chunk = []
        for child in children:
            if child.get('type') == 'inline_html':
                raw = child.get('raw') or child.get('text', '')
                break_match = _BREAK_HTML_RE.match(raw)
                if break_match:
                    if chunk:
                        last_paragraph = self._build_paragraph_from_children(chunk, style_name)
                        chunk = []
                    self._handle_break_marker(f"<!--break:{break_match.group(1)}-->")
                    self._suppress_next_first_line_indent = True
                    continue
            chunk.append(child)

        if chunk:
            last_paragraph = self._build_paragraph_from_children(chunk, style_name)

        return last_paragraph

    def _get_infobox_style_case_insensitive(self, key: str):
        key_lower = key.lower()
        for k, v in self.config.infobox_styles.items():
            if k.lower() == key_lower:
                return v
        return None

    def _get_single_style_case_insensitive(self, key: str):
        key_lower = key.lower()
        for k, v in self.config.single_paragraph_styles.items():
            if k.lower() == key_lower:
                return v
        return None

    def _parse_inline_children(self, text: str):
        if not text:
            return []
        parser = self._markdown_parser or mistune.create_markdown(renderer=None, plugins=['table'])
        ast = parser(text)
        if ast and ast[0].get('type') == 'paragraph':
            return ast[0].get('children', [])
        return [{'type': 'text', 'raw': text}]

    def _handle_style_marker(self, node: dict, full_text: str):
        if full_text.startswith("<!--block_start:"):
            match = re.match(r"<!--block_start:([^>]+)-->", full_text)
            if match:
                key = match.group(1).strip()
                style_config = self._get_infobox_style_case_insensitive(key)
                if style_config:
                    self._current_block_style = style_config.style_name
                    self._current_block_style_config = style_config
                    self._is_first_para_in_block = True
        elif full_text.startswith("<!--block_end:"):
            self._current_block_style = None
            self._current_block_style_config = None
        elif full_text.startswith("<!--single_tag:"):
            match = re.match(r"<!--single_tag:([^:]+):(.*)-->", full_text, re.DOTALL)
            if match:
                key = match.group(1).strip()
                content = match.group(2).strip()
                style_info = self._get_single_style_case_insensitive(key)
                if style_info:
                    p = self.doc.add_paragraph()
                    try:
                        p.style = style_info['style_name']
                        self._flatten_style_properties(p)
                    except Exception:
                        pass
                    _add_formatted_text(
                        p,
                        self._parse_inline_children(content),
                        part=self.doc.part,
                        link_style=self._hyperlink_style_id
                    )
                    self._apply_pending_break_indent_rule(p)
                    if self._suppress_next_first_line_indent_from_arrow:
                        p.paragraph_format.first_line_indent = Pt(0)
                        self._suppress_next_first_line_indent_from_arrow = False
                    self._apply_indent_from_arrows(p)
                    return p
                else:
                    style_config = self._get_infobox_style_case_insensitive(key)
                    if style_config:
                        p = self.doc.add_paragraph()
                        try:
                            p.style = style_config.style_name
                            self._flatten_style_properties(p)
                        except Exception:
                            pass
                        _add_formatted_text(
                            p,
                            self._parse_inline_children(content),
                            part=self.doc.part,
                            link_style=self._hyperlink_style_id
                        )
                        self._apply_pending_break_indent_rule(p)
                        if self._suppress_next_first_line_indent_from_arrow:
                            p.paragraph_format.first_line_indent = Pt(0)
                            self._suppress_next_first_line_indent_from_arrow = False
                        self._apply_indent_from_arrows(p)
                        return p

    def _handle_break_marker(self, full_text: str):
        match = re.match(r"<!--break:(page|column)-->", full_text)
        if not match:
            return None

        p = self.doc.add_paragraph()
        run = p.add_run()
        if match.group(1) == 'page':
            run.add_break(WD_BREAK.PAGE)
        else:
            run.add_break(WD_BREAK.COLUMN)
        return p

    def _get_line_style_key(self, full_text: str) -> str | None:
        match = re.match(r"<!--line_style:([^>|]+?)(?:-->|\|)", full_text)
        if match:
            return match.group(1).strip()
        return None

    def _get_line_style_arrows(self, full_text: str) -> str:
        match = re.search(r"\|arrows:([›→⋙]+)-->", full_text)
        if match:
            return match.group(1)
        return ""

    def _get_line_style_name(self, key: str) -> str:
        if not key:
            return "Horizontal Line"
        style = self.config.get_line_style(key)
        if style:
            return style.style_name
        return "Horizontal Line"
