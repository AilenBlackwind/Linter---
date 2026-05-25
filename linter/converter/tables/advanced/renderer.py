from typing import Dict, Any, List, Optional, Tuple
from docx.table import Table, _Cell
from docx.shared import Pt

from linter.converter.tables.advanced.models import (
    TableStyle,
    RowTypeStyle,
    CellOverride,
    ColorRule,
    BorderStyle,
    CellMargins,
    BorderSide,
)
from linter.converter.tables.advanced.rules_engine import RulesEngine
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def cell_ref_to_indices(cell_ref: str) -> Tuple[int, int]:
    cell_ref = cell_ref.strip().upper()

    col_str = ""
    row_str = ""

    for char in cell_ref:
        if char.isalpha():
            col_str += char
        else:
            row_str += char

    col = 0
    for char in col_str:
        col = col * 26 + (ord(char) - ord('A') + 1)
    col -= 1

    row = int(row_str) - 1

    return (row, col)


def indices_to_cell_ref(row_idx: int, col_idx: int) -> str:
    col = col_idx
    col_str = ""
    while True:
        col_str = chr(ord('A') + (col % 26)) + col_str
        col = col // 26 - 1
        if col < 0:
            break

    row_str = str(row_idx + 1)

    return col_str + row_str


def _get_or_create_tcPr(cell: _Cell):
    return cell._tc.get_or_add_tcPr()


def _get_or_create_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def _apply_cell_shading(cell: _Cell, fill_color: str):
    tcPr = _get_or_create_tcPr(cell)
    shd = _get_or_create_child(tcPr, 'w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color.upper())


def _clear_cell_shading(cell: _Cell):
    tcPr = _get_or_create_tcPr(cell)
    shd = _get_or_create_child(tcPr, 'w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF')


def _apply_cell_border(cell: _Cell, side: BorderSide, border_style: BorderStyle):
    tcPr = _get_or_create_tcPr(cell)
    tcBorders = _get_or_create_child(tcPr, 'w:tcBorders')

    side_map = {
        'top': 'w:top',
        'bottom': 'w:bottom',
        'left': 'w:left',
        'right': 'w:right',
    }

    if side == 'all':
        for s in ['top', 'bottom', 'left', 'right']:
            _apply_cell_border(cell, s, border_style)
        return

    if side == 'horizontal':
        _apply_cell_border(cell, 'top', border_style)
        _apply_cell_border(cell, 'bottom', border_style)
        return

    if side == 'vertical':
        _apply_cell_border(cell, 'left', border_style)
        _apply_cell_border(cell, 'right', border_style)
        return

    if side in side_map:
        tag = side_map[side]
        border_el = _get_or_create_child(tcBorders, tag)

        if border_style.val == 'nil' or border_style.size <= 0:
            border_el.set(qn('w:val'), 'nil')
        else:
            border_el.set(qn('w:val'), border_style.val)
            border_el.set(qn('w:sz'), str(border_style.size))
            border_el.set(qn('w:color'), border_style.color.upper())


def _apply_cell_margins(cell: _Cell, margins: CellMargins):
    tcPr = _get_or_create_tcPr(cell)
    tcMar = _get_or_create_child(tcPr, 'w:tcMar')

    for margin_name, tag_name in [
        ('top', 'w:top'),
        ('start', 'w:left'),
        ('bottom', 'w:bottom'),
        ('end', 'w:right'),
    ]:
        value = getattr(margins, margin_name)
        el = _get_or_create_child(tcMar, tag_name)
        el.set(qn('w:w'), str(value))
        el.set(qn('w:type'), 'dxa')


def _apply_text_style_to_cell(cell: _Cell, font_color: Optional[str] = None,
                              bold: Optional[bool] = None,
                              italic: Optional[bool] = None,
                              font_size: Optional[int] = None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if font_color is not None:
                try:
                    run.font.color.rgb = font_color
                except:
                    pass
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size is not None:
                run.font.size = Pt(font_size)

        if not paragraph.runs and paragraph.text:
            text = paragraph.text
            paragraph.clear()
            run = paragraph.add_run(text)
            if font_color is not None:
                try:
                    run.font.color.rgb = font_color
                except:
                    pass
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size is not None:
                run.font.size = Pt(font_size)


class AdvancedTableRenderer:

    def __init__(self, table_style: TableStyle):
        self.style = table_style
        self.rules_engine = RulesEngine(table_style.color_rules)

        self._cell_override_cache: Dict[Tuple[int, int], CellOverride] = {}
        for override in table_style.cell_overrides:
            try:
                row_idx, col_idx = cell_ref_to_indices(override.cell_ref)
                self._cell_override_cache[(row_idx, col_idx)] = override
            except:
                pass

    def _get_cell_style(self, row_idx: int, col_idx: int, cell_text: str,
                         row_count: int, col_count: int) -> Dict[str, Any]:
        style: Dict[str, Any] = {}

        if self.style.default_shading:
            style['shading'] = self.style.default_shading
        if self.style.default_borders:
            style['borders'] = dict(self.style.default_borders)
        style['cell_margins'] = self.style.cell_margins

        row_type_style: Optional[RowTypeStyle] = None

        if row_idx == 0:
            if 'header' in self.style.row_types:
                row_type_style = self.style.row_types['header']
        elif row_idx == row_count - 1:
            if 'last_row' in self.style.row_types:
                row_type_style = self.style.row_types['last_row']
        elif row_idx % 2 == 0:
            if 'odd' in self.style.row_types:
                row_type_style = self.style.row_types['odd']
        else:
            if 'even' in self.style.row_types:
                row_type_style = self.style.row_types['even']

        if row_type_style:
            if row_type_style.shading:
                style['shading'] = row_type_style.shading
            if row_type_style.font_color:
                style['font_color'] = row_type_style.font_color
            if row_type_style.bold is not None:
                style['bold'] = row_type_style.bold
            if row_type_style.italic is not None:
                style['italic'] = row_type_style.italic
            if row_type_style.borders:
                if 'borders' not in style:
                    style['borders'] = {}
                for side, bs in row_type_style.borders.items():
                    style['borders'][side] = bs

        col_type_style: Optional[RowTypeStyle] = None

        if col_idx == 0:
            if 'first_column' in self.style.row_types:
                col_type_style = self.style.row_types['first_column']
        elif col_idx == col_count - 1:
            if 'last_column' in self.style.row_types:
                col_type_style = self.style.row_types['last_column']

        if col_type_style:
            if col_type_style.shading:
                style['shading'] = col_type_style.shading
            if col_type_style.font_color:
                style['font_color'] = col_type_style.font_color
            if col_type_style.bold is not None:
                style['bold'] = col_type_style.bold
            if col_type_style.italic is not None:
                style['italic'] = col_type_style.italic
            if col_type_style.borders:
                if 'borders' not in style:
                    style['borders'] = {}
                for side, bs in col_type_style.borders.items():
                    style['borders'][side] = bs

        cell_key = (row_idx, col_idx)
        if cell_key in self._cell_override_cache:
            override = self._cell_override_cache[cell_key]
            if override.shading:
                style['shading'] = override.shading
            if override.font_color:
                style['font_color'] = override.font_color
            if override.bold is not None:
                style['bold'] = override.bold
            if override.italic is not None:
                style['italic'] = override.italic
            if override.borders:
                if 'borders' not in style:
                    style['borders'] = {}
                for side, bs in override.borders.items():
                    style['borders'][side] = bs

        style = self.rules_engine.apply_rules(cell_text, style, col_idx=col_idx)

        return style

    def _apply_style_to_cell(self, cell: _Cell, style: Dict[str, Any]):
        if 'cell_margins' in style:
            _apply_cell_margins(cell, style['cell_margins'])

        if 'shading' in style and style['shading']:
            _apply_cell_shading(cell, style['shading'])

        if 'borders' in style and style['borders']:
            for side, border_style in style['borders'].items():
                _apply_cell_border(cell, side, border_style)

        font_color = style.get('font_color')
        bold = style.get('bold')
        italic = style.get('italic')

        if font_color is not None or bold is not None or italic is not None:
            _apply_text_style_to_cell(
                cell,
                font_color=font_color,
                bold=bold,
                italic=italic,
            )

    def apply_to_table(self, table: Table):
        row_count = len(table.rows)
        if row_count == 0:
            return

        col_count = len(table.rows[0].cells)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = ""
                for paragraph in cell.paragraphs:
                    cell_text += paragraph.text + " "
                cell_text = cell_text.strip()

                style = self._get_cell_style(
                    row_idx, col_idx, cell_text, row_count, col_count
                )

                self._apply_style_to_cell(cell, style)

        from linter.converter.tables.processor import apply_text_wrap
        apply_text_wrap(table, self.style.text_wrap,
                        self.style.left_from_text, self.style.right_from_text)

        if 'header' in self.style.row_types:
            header_style = self.style.row_types['header']
            if header_style.repeat and row_count > 0:
                try:
                    from linter.converter.tables.processor import set_repeat_table_header as set_repeat
                    first_row = table.rows[0]
                    trPr = first_row._tr.get_or_add_trPr()
                    tblHeader = trPr.find(qn('w:tblHeader'))
                    if tblHeader is None:
                        tblHeader = OxmlElement('w:tblHeader')
                        trPr.append(tblHeader)
                    tblHeader.set(qn('w:val'), '1')
                except:
                    pass

    @classmethod
    def from_dict(cls, style_data: dict, name: str = "") -> 'AdvancedTableRenderer':
        table_style = TableStyle.from_dict(style_data, name)
        return cls(table_style)
